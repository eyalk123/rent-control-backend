"""Unit tests for DocumentExtractionService (Anthropic client mocked)."""
import base64
import io

import pytest
from fastapi import HTTPException

from app.schemas.document_extraction import (
    ExtractedRenter,
    FieldNote,
    LeaseExtraction,
    LeaseYearGuess,
)
from app.services.document_extraction_service import (
    DocumentExtractionService,
    _clean_extraction,
)

PDF_MEDIA = "application/pdf"
DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _service(api_key="test-key"):
    return DocumentExtractionService(api_key=api_key, model="claude-opus-4-8")


def _one_renter() -> LeaseExtraction:
    """A LeaseExtraction seeded with a single renter (the common single-tenant case)."""
    return LeaseExtraction(renters=[ExtractedRenter()])


class _FakeMessages:
    def __init__(self, response):
        self._response = response
        self.calls = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        return self._response


class _FakeClient:
    def __init__(self, response):
        self.messages = _FakeMessages(response)


class _FakeUsage:
    input_tokens = 1000
    output_tokens = 500
    cache_read_input_tokens = 0
    cache_creation_input_tokens = 0


class _FakeToolUse:
    type = "tool_use"
    name = "record_lease_extraction"

    def __init__(self, tool_input):
        self.input = tool_input


class _FakeResponse:
    def __init__(self, content, stop_reason="tool_use"):
        self.content = content
        self.stop_reason = stop_reason
        self.usage = _FakeUsage()


# --- content block construction (per file type) ---

def _make_pdf(pages_text: list[str]) -> bytes:
    """Build a small real PDF (one page per string) with fpdf2 for rasterization tests."""
    from fpdf import FPDF

    pdf = FPDF()
    for text in pages_text:
        pdf.add_page()
        pdf.set_font("Helvetica", size=16)
        pdf.cell(0, 10, text)
    return bytes(pdf.output())


def test_pdf_renders_to_image_blocks():
    blocks = _service()._build_content_blocks(_make_pdf(["Lease 13 Ein Harod"]), PDF_MEDIA)
    assert len(blocks) == 1
    assert blocks[0]["type"] == "image"
    assert blocks[0]["source"]["media_type"] == "image/png"
    # The rendered page is a real PNG (magic bytes), not the raw PDF text layer.
    assert base64.standard_b64decode(blocks[0]["source"]["data"]).startswith(b"\x89PNG")


def test_pdf_renders_one_image_block_per_page():
    blocks = _service()._build_content_blocks(_make_pdf(["Page one", "Page two"]), PDF_MEDIA)
    assert len(blocks) == 2
    assert all(b["type"] == "image" for b in blocks)


def test_corrupt_pdf_bytes_raise_422():
    with pytest.raises(HTTPException) as exc:
        _service()._build_content_blocks(b"not a pdf", PDF_MEDIA)
    assert exc.value.status_code == 422


@pytest.mark.parametrize("media", ["image/png", "image/jpeg", "image/webp", "image/gif"])
def test_image_builds_image_block(media):
    blocks = _service()._build_content_blocks(b"\x89PNG bytes", media)
    assert len(blocks) == 1
    assert blocks[0]["type"] == "image"
    assert blocks[0]["source"]["media_type"] == media
    assert base64.standard_b64decode(blocks[0]["source"]["data"]) == b"\x89PNG bytes"


def test_content_type_with_charset_suffix_is_handled():
    blocks = _service()._build_content_blocks(
        _make_pdf(["Lease"]), "application/pdf; charset=binary"
    )
    assert blocks[0]["type"] == "image"


def test_docx_builds_text_block_from_paragraphs_and_tables():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Lease for 12 Herzl St, Tel Aviv")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Rent"
    table.rows[0].cells[1].text = "5000"
    buf = io.BytesIO()
    doc.save(buf)

    blocks = _service()._build_content_blocks(buf.getvalue(), DOCX_MEDIA)
    assert len(blocks) == 1
    assert blocks[0]["type"] == "text"
    assert "12 Herzl St" in blocks[0]["text"]
    assert "Rent | 5000" in blocks[0]["text"]


def test_unsupported_type_raises_415():
    with pytest.raises(HTTPException) as exc:
        _service()._build_content_blocks(b"x", "text/csv")
    assert exc.value.status_code == 415


# --- extract_lease orchestration ---

def test_extract_lease_returns_result_with_meta(monkeypatch):
    svc = _service()
    draft = _one_renter()
    draft.property.city = "Tel Aviv"
    draft.renters[0].base_rent = 5000.0
    draft.notes = [FieldNote(section="renter", field="base_rent", renter_index=0, confidence="medium", source_text="5,000 ₪")]
    fake = _FakeClient(_FakeResponse([_FakeToolUse(draft.model_dump())]))
    monkeypatch.setattr(svc, "_client", lambda: fake)

    result = svc.extract_lease(_make_pdf(["Lease"]), PDF_MEDIA)

    assert result.extraction.property.city == "Tel Aviv"
    assert result.extraction.renters[0].base_rent == 5000.0
    # Telemetry computed from the response + parsed tool input.
    assert result.meta.input_tokens == 1000
    assert result.meta.output_tokens == 500
    assert result.meta.fields_extracted == 2  # city + base_rent populated
    assert result.meta.medium_confidence_count == 1  # from the note
    assert result.meta.estimated_cost_usd is not None
    # A non-strict extraction tool was forced, and the PDF was sent as a rendered image
    # (not a native document block that would carry the corrupt text layer).
    call = fake.messages.calls[0]
    assert call["tool_choice"]["name"] == "record_lease_extraction"
    assert call["messages"][0]["content"][0]["type"] == "image"


def test_extract_lease_refusal_raises_422(monkeypatch):
    svc = _service()
    fake = _FakeClient(_FakeResponse([], stop_reason="refusal"))
    monkeypatch.setattr(svc, "_client", lambda: fake)
    with pytest.raises(HTTPException) as exc:
        svc.extract_lease(_make_pdf(["Lease"]), PDF_MEDIA)
    assert exc.value.status_code == 422


def test_extract_lease_no_tool_call_raises_422(monkeypatch):
    svc = _service()
    fake = _FakeClient(_FakeResponse([], stop_reason="max_tokens"))
    monkeypatch.setattr(svc, "_client", lambda: fake)
    with pytest.raises(HTTPException) as exc:
        svc.extract_lease(_make_pdf(["Lease"]), PDF_MEDIA)
    assert exc.value.status_code == 422


def test_missing_api_key_raises_503():
    with pytest.raises(HTTPException) as exc:
        _service(api_key="").extract_lease(_make_pdf(["Lease"]), PDF_MEDIA)
    assert exc.value.status_code == 503


# --- validation cleaning (don't hand the user values that would error on submit) ---

def test_clean_drops_out_of_range_payment_day():
    e = _one_renter()
    e.renters[0].payment_day_of_month = 45
    _clean_extraction(e)
    assert e.renters[0].payment_day_of_month is None


def test_clean_keeps_valid_payment_day():
    e = _one_renter()
    e.renters[0].payment_day_of_month = 10
    _clean_extraction(e)
    assert e.renters[0].payment_day_of_month == 10


def test_clean_drops_non_iso_lease_start():
    e = _one_renter()
    e.renters[0].lease_start = "see addendum"
    _clean_extraction(e)
    assert e.renters[0].lease_start is None


def test_clean_keeps_iso_lease_start():
    e = _one_renter()
    e.renters[0].lease_start = "2025-03-01"
    _clean_extraction(e)
    assert e.renters[0].lease_start == "2025-03-01"


def test_clean_drops_negative_amounts():
    e = _one_renter()
    e.renters[0].base_rent = -500.0
    e.property.number_of_rooms = -3
    e.renters[0].lease_years = [LeaseYearGuess(amount=-1, type="contract")]
    _clean_extraction(e)
    assert e.renters[0].base_rent is None
    assert e.property.number_of_rooms is None
    assert e.renters[0].lease_years is None


def test_clean_drops_note_for_nulled_field():
    e = _one_renter()
    e.renters[0].payment_day_of_month = 45  # invalid → will be nulled
    e.renters[0].base_rent = 5000.0
    e.notes = [
        FieldNote(section="renter", field="payment_day_of_month", renter_index=0, confidence="low", source_text="?"),
        FieldNote(section="renter", field="base_rent", renter_index=0, confidence="medium", source_text="5,000"),
    ]
    _clean_extraction(e)
    # The note for the nulled field is removed; the one for a surviving field stays.
    assert [n.field for n in e.notes] == ["base_rent"]


def test_clean_drops_unknown_insurance_type():
    e = _one_renter()
    e.renters[0].insurance_type = "property insurance"  # not one of the enum values
    _clean_extraction(e)
    assert e.renters[0].insurance_type is None


def test_clean_keeps_known_insurance_type():
    e = _one_renter()
    e.renters[0].insurance_type = "bank_guarantee"
    _clean_extraction(e)
    assert e.renters[0].insurance_type == "bank_guarantee"


def test_clean_drops_national_id_mistaken_for_phone():
    e = _one_renter()
    e.renters[0].phone = "312345678"  # bare 9-digit national ID, not a phone
    _clean_extraction(e)
    assert e.renters[0].phone is None


def test_clean_keeps_real_phone():
    e = _one_renter()
    e.renters[0].phone = "0521234567"  # 10-digit Israeli mobile
    _clean_extraction(e)
    assert e.renters[0].phone == "0521234567"


def test_clean_drops_negative_joint_rent():
    e = LeaseExtraction(renters=[ExtractedRenter(), ExtractedRenter()], rent_is_joint=True)
    e.joint_monthly_rent = -100.0
    _clean_extraction(e)
    assert e.joint_monthly_rent is None


def test_field_stats_counts_all_renters():
    e = LeaseExtraction(
        renters=[ExtractedRenter(first_name="Dana"), ExtractedRenter(first_name="Noa")],
    )
    e.property.city = "Tel Aviv"
    result = _clean_extraction(e)
    from app.services.document_extraction_service import _field_stats

    extracted, _, _ = _field_stats(result)
    assert extracted == 3  # city + two first_names


def test_clean_scopes_notes_to_the_right_renter():
    e = LeaseExtraction(renters=[ExtractedRenter(), ExtractedRenter(base_rent=4000.0)])
    e.notes = [
        # points at renter 0 whose base_rent is null → dropped
        FieldNote(section="renter", field="base_rent", renter_index=0, confidence="low", source_text="?"),
        # points at renter 1 whose base_rent is set → kept
        FieldNote(section="renter", field="base_rent", renter_index=1, confidence="medium", source_text="4,000"),
    ]
    _clean_extraction(e)
    assert [(n.field, n.renter_index) for n in e.notes] == [("base_rent", 1)]
