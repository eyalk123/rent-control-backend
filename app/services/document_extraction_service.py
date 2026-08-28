"""Lease document extraction via Claude (vision + structured output).

The service is the only place that talks to the Anthropic API. It receives the
raw upload bytes, builds the right content blocks for the file type, and asks Claude
to return a :class:`LeaseExtraction` conforming to our schema.

PDFs are rasterized to page images (via PyMuPDF) and sent as `image` blocks rather
than as a native `document` block. Some leases embed subsetted fonts with a broken
ToUnicode map, so the PDF's hidden text layer reports wrong digits (rent, dates, unit
numbers) even though the page *renders* correctly. Anthropic's native PDF handling
feeds both the rendered pixels and that corrupt text layer to the model, and the model
tends to trust the authoritative-looking text. Rendering to pixels ourselves drops the
text layer entirely, so Claude reads only what a human sees. Images go to Claude
directly; DOCX (real text, no glyph problem) is converted to text with python-docx.

The file is processed in-memory and never written to disk or Firebase — clients hold
the original and attach it to Firebase only when the reviewed form is submitted.
"""
import base64
import io
import logging
from dataclasses import dataclass
from datetime import date
from typing import Optional

from anthropic import Anthropic
from fastapi import HTTPException, status
from pydantic import ValidationError

from app.schemas.document_extraction import (
    ExtractedProperty,
    ExtractedRenter,
    LeaseExtraction,
)

logger = logging.getLogger(__name__)

# A non-strict tool the model fills with the extracted data. We deliberately avoid
# structured outputs (messages.parse): the strict grammar compiler rejects a schema
# this large. The schema is still used to validate the tool's output afterwards.
_TOOL_NAME = "record_lease_extraction"
_EXTRACTION_TOOL = {
    "name": _TOOL_NAME,
    "description": "Record the structured property and renter data extracted from the lease document.",
    "input_schema": LeaseExtraction.model_json_schema(),
}

# USD per 1M tokens (input, output). Used for a rough cost estimate on the audit log.
_PRICES: dict[str, tuple[float, float]] = {
    "claude-sonnet-4-6": (3.0, 15.0),
    "claude-opus-4-8": (5.0, 25.0),
}


@dataclass
class ExtractionMeta:
    """Telemetry about one extraction call, stored on the audit log."""

    model: str
    input_tokens: Optional[int]
    output_tokens: Optional[int]
    cache_read_tokens: Optional[int]
    cache_creation_tokens: Optional[int]
    estimated_cost_usd: Optional[float]
    fields_extracted: int
    low_confidence_count: int
    medium_confidence_count: int


@dataclass
class ExtractionResult:
    extraction: LeaseExtraction
    meta: ExtractionMeta

# MIME types Claude reads natively as images.
_IMAGE_MEDIA_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
_PDF_MEDIA_TYPE = "application/pdf"
_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# PDF page rendering. 200 DPI keeps small digits and dense tables sharp for the vision
# model (the only real precision lever when reading from pixels). The page cap is a
# defensive guard against pathological uploads — leases are a handful of pages.
_RENDER_DPI = 200
_MAX_PDF_PAGES = 20

# One descriptive line per field — the "field schema" Claude maps the document onto.
# Kept stable so it can be prompt-cached across requests.
#
# DISABLED extraction fields — rarely stated in a lease, so we don't spend tokens extracting
# them. To re-enable one: uncomment its field in app/schemas/document_extraction.py AND paste
# its bullet back into the matching (Property / Renter) section of _SYSTEM_PROMPT below.
#   Property:
#   - zip_code: the property's postal code (part of the address/city/zip group).
#   - sq_ft: floor area as a number (whatever unit the document uses).
#   - parking_numbers: list of parking spot identifiers.
#   - block, plot: land registry block ("gush") / plot ("helka"). (Restore alongside `apartment`.)
#   - electricity_meter_number, electricity_account_number, water_meter_number, water_account_number: utility identifiers.
#   - property_tax, house_committee: periodic property tax ("arnona") and building-committee ("vaad bayit") amounts.
#   - inventory_notes: any inventory / contents description.
_SYSTEM_PROMPT = """You extract structured data from rental lease / property contracts to pre-fill a property-management app's forms. Documents are often in Hebrew (right-to-left) and may mix Hebrew, English, and numbers in tables — read them carefully and preserve the correct values.

A single lease usually describes a property and one or more renters (co-tenants who sign the same lease). Populate the property once and add one entry to `renters` for EACH tenant. Fill EVERY field that the document states — a typical lease contains most of them. Leave a field null ONLY if the document genuinely doesn't contain it; never guess or invent values. Dates as ISO YYYY-MM-DD; money/areas as plain numbers (no currency symbols or commas).

Confidence: for any field you are NOT highly confident about (inferred, ambiguous, or loosely derived), append one entry to `notes` with its `section` ("property" or "renter"), `field` (the exact field name below), `renter_index` (for a renter field, the 0-based index of the renter in `renters`; null for a property field), `confidence` ("medium" or "low"), and `source_text` (the short verbatim snippet it came from). Do NOT add a note for a field you're confident about, and never add a note for a null field.

Property fields:
- address_evidence: FILL THIS FIRST, before `address`. Copy the short verbatim phrase from the document that states WHERE THE RENTED PROPERTY IS — the clause describing the המושכר / הדירה / הנכס (e.g. "הדירה ברחוב הרצל 12 תל אביב"). Then derive `address`/`city` from this same phrase. Null only if the document truly never states the property's location.
- address, city: the address of the PROPERTY BEING RENTED (the demised premises) — NEVER a person's address. Read carefully; this is the most important field:
  * The rented property is called המושכר / הדירה / הנכס / הממכר. Its address is stated where the lease describes WHAT is being rented — usually the opening recitals ("הואיל והמשכיר הינו בעל הזכויות בדירה ברחוב ___ בעיר ___") or a dedicated clause ("המושכר: דירה ברחוב ___") or the "הואיל" / "מבוא" section. Take `address` and `city` from there.
  * Do NOT use any address from the parties block / signature area (the "בין ___ לבין ___" part). There each party's name is followed by "ת.ז ___" and "מרחוב ___" / "מ___" — those are the LANDLORD's and the TENANT's OWN home/mailing addresses. They are commonly the FIRST addresses in the document; taking one of them by mistake is the single most common error — do not.
  * If several street addresses appear, choose the one tied to המושכר/הדירה/הנכס, not the one tied to a person's name or ת.ז. If still unsure, pick the most likely property address and add a `notes` entry for `address` with its source_text.
  * `address` is the STREET NAME AND HOUSE NUMBER ONLY — never the apartment/unit number, and never the floor. The apartment number, when present, goes in `apartment` (below), NOT in `address`.
- type: one of apartment, house, commercial, garden_apartment, housing_unit.
- number_of_rooms, floor: room count and floor number.
- apartment: the apartment/unit number. In Hebrew leases this is very often written right after the street, frequently in parentheses — e.g. "ברח' בלפור 51 (דירה מס' 3)" → `address` = "בלפור 51", `apartment` = "3". Also handles forms like "דירה 3", "דירה מס' 3", "apt 3", "יח' 3". Extract just the number/identifier into `apartment` and keep it OUT of `address`.
- property_owner: the landlord/owner's name.

Renter fields — one `renters` entry PER tenant. If several people sign as tenants, list them all; do not merge them or push them into extra_contacts.
- first_name, last_name: the tenant's given and family name.
- phone: the tenant's PHONE number only. An Israeli phone is typically 9-10 digits starting with 0 (mobile 05X-XXXXXXX) or +972. Before filling this, check the number really looks like a phone. A 9-digit national ID number (תעודת זהות / ת"ז) is NOT a phone — if the only number you see near the tenant is an ID, leave phone null rather than putting the ID here.
- email: the tenant's EMAIL address, exactly as written (one "@", no spaces). Take it only from the tenant's own details — an email in the landlord's / agent's / lawyer's block belongs to none of these fields, so leave this null rather than borrowing one.
- lease_start: the lease commencement date (ISO YYYY-MM-DD).
- payment_type: how the rent is paid — exactly one of "bank_transfer" (העברה בנקאית / הוראת קבע), "check" (המחאות / צ'קים / שיקים), "cash" (מזומן) or "bit" (ביט). Return null when the document doesn't say, or describes a method that is none of these — do not stretch one to fit.
- payment_day_of_month: the day OF THE MONTH rent is due — a single whole number 1-31, nothing else. It is NOT a phone number, NOT a 9-digit ID (ת"ז), NOT a bank account/branch number, NOT a full date, NOT a sum of money. In Hebrew it reads like "בכל 1 לחודש", "עד ה-10 בכל חודש", "ב-1 לכל חודש" — take just the day number from such a clause. If the document never says which day of the month rent is due, return null; never borrow a nearby number just because one appears next to the payment terms.
- insurance_type: the required security/collateral type — exactly one of "bank_guarantee" (ערבות בנקאית) or "wire_transfer" (פיקדון / העברה בנקאית / cash deposit). insurance_amount: its amount (usually stated right next to the type). Return null for insurance_type if it isn't one of those two.
- number_of_payments: installments per year (e.g. 12 for monthly).
- extra_contacts: NON-tenant contacts only, e.g. guarantors (ערבים), each {name, phone}. People who sign as tenants belong in `renters`, not here.

Joint vs. per-tenant rent: most leases with several tenants state ONE joint rent for all of them together (they are jointly liable, "ביחד ולחוד"), not a separate amount per tenant. When the rent is joint, set `rent_is_joint` true, put that single first-year MONTHLY total in top-level `joint_monthly_rent`, and leave every renter's `base_rent` null (still fill the shared lease-term fields — escalation, years, insurance, payment — identically on each renter). Only when the document gives each tenant their OWN distinct amount, set `rent_is_joint` false and fill each renter's own `base_rent`. For a single-tenant lease, set `rent_is_joint` false and fill that renter's `base_rent`.

Lease term — describe it as INTENT; the app rebuilds the year-by-year schedule from these:
- contract_term_years: number of binding (contract) years. option_years: number of renewal-option years.
- contract_term_months / option_term_months: any ODD MONTHS ON TOP of those whole years, 0-11. A term stated as "24 months" is contract_term_years 2 and contract_term_months 0; "28 months" or "two years and four months" is contract_term_years 2 and contract_term_months 4; "18 months" is 1 and 6; a term shorter than a year ("8 months") is contract_term_years 0 and contract_term_months 8. Return 0 (not null) when the term is a whole number of years.
- base_rent: the FIRST-YEAR MONTHLY rent (a single monthly figure, never annual). Leave null when rent_is_joint is true (use joint_monthly_rent instead).
- rent_escalation_mode: how the monthly rent changes each year — "none" (flat, same every year), "percent" (rises a fixed % each year), "fixed" (rises a fixed money amount each year), "cpi" (linked/indexed to the Consumer Price Index — מדד המחירים לצרכן / הצמדה למדד), or "custom" (irregular per-year amounts that follow no single rule). Use "cpi" whenever the rent is tied to the index (הצמדה למדד המחירים לצרכן), even if a minimum increase is also mentioned.
- rent_escalation_value: the percent (for "percent") or the money amount (for "fixed"). Null for "none"/"custom"/"cpi".
- lease_years: leave EMPTY unless rent_escalation_mode is "custom". When custom, list one row per lease year, each {amount: that year's MONTHLY rent in the SAME unit as base_rent, type: "contract" or "option"}, and set contract_term_years + option_years to equal the number of rows.

Return only the structured data."""


class DocumentExtractionService:
    def __init__(self, api_key: str, model: str):
        self._api_key = api_key
        self._model = model

    @property
    def model_name(self) -> str:
        return self._model

    def _client(self) -> Anthropic:
        if not self._api_key:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Document extraction is not configured (missing ANTHROPIC_API_KEY).",
            )
        return Anthropic(api_key=self._api_key)

    def _build_content_blocks(self, file_bytes: bytes, content_type: str) -> list[dict]:
        """Turn the upload into a list of Claude content blocks, by file type."""
        media_type = (content_type or "").split(";")[0].strip().lower()

        if media_type == _PDF_MEDIA_TYPE:
            # Rasterize to page images so the corrupt PDF text layer never reaches the
            # model (see the module docstring).
            return self._pdf_to_image_blocks(file_bytes)

        if media_type in _IMAGE_MEDIA_TYPES:
            return [{
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": base64.standard_b64encode(file_bytes).decode(),
                },
            }]

        if media_type == _DOCX_MEDIA_TYPE:
            text = self._docx_to_text(file_bytes)
            if not text.strip():
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="The document appears to be empty.",
                )
            return [{"type": "text", "text": text}]

        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unsupported file type. Upload a PDF, DOCX, or image (JPEG/PNG/GIF/WebP).",
        )

    @staticmethod
    def _pdf_to_image_blocks(file_bytes: bytes) -> list[dict]:
        """Render each PDF page to a PNG image block. Bypasses the (sometimes corrupt)
        embedded text layer by reading only the rendered pixels — see module docstring."""
        import fitz  # PyMuPDF, imported lazily so the dep is only needed at runtime

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="The document could not be processed.",
            )

        zoom = _RENDER_DPI / 72
        matrix = fitz.Matrix(zoom, zoom)
        blocks: list[dict] = []
        try:
            for page in doc:
                if len(blocks) >= _MAX_PDF_PAGES:
                    break
                pixmap = page.get_pixmap(matrix=matrix)
                blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/png",
                        "data": base64.standard_b64encode(pixmap.tobytes("png")).decode(),
                    },
                })
        finally:
            doc.close()

        if not blocks:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="The document appears to be empty.",
            )
        return blocks

    @staticmethod
    def _docx_to_text(file_bytes: bytes) -> str:
        """Extract paragraph and table text from a .docx, preserving reading order."""
        from docx import Document  # imported lazily so the dep is only needed at runtime

        document = Document(io.BytesIO(file_bytes))
        lines: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    def extract_lease(self, file_bytes: bytes, content_type: str) -> ExtractionResult:
        """Extract a property + renter draft from a lease document, with call telemetry."""
        content_blocks = self._build_content_blocks(file_bytes, content_type)
        client = self._client()

        # Use a NON-strict tool (not structured-outputs / messages.parse): the strict
        # grammar compiler rejects a schema this large ("compiled grammar is too large").
        # The model fills the tool's schema best-effort and we validate it ourselves.
        response = client.messages.create(
            model=self._model,
            max_tokens=8192,
            system=[
                {
                    "type": "text",
                    "text": _SYSTEM_PROMPT,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            tools=[_EXTRACTION_TOOL],
            tool_choice={"type": "tool", "name": _TOOL_NAME},
            messages=[
                {
                    "role": "user",
                    "content": [
                        *content_blocks,
                        {
                            "type": "text",
                            "text": "Extract the property and renter details from this lease document.",
                        },
                    ],
                }
            ],
        )

        if response.stop_reason == "refusal":
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="The document could not be processed.",
            )
        tool_block = next(
            (b for b in response.content if getattr(b, "type", None) == "tool_use"), None
        )
        if tool_block is None:
            # e.g. stop_reason == "max_tokens" — no complete tool call returned.
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not extract structured data from the document.",
            )
        try:
            parsed = LeaseExtraction.model_validate(tool_block.input)
        except ValidationError:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not extract structured data from the document.",
            )

        # Only logged when the model emitted a value we had to throw away — so a bogus
        # field (e.g. a phone number as payment_day_of_month) is provable after the fact
        # instead of silently vanishing. Deliberately narrow: we log ONLY the discarded
        # values, never the whole extraction, which is full of tenant PII.
        discarded: list[str] = []
        extraction = _clean_extraction(parsed, discarded)
        if discarded:
            logger.warning(
                "Lease extraction (model=%s) returned unusable values, discarded: %s",
                self._model,
                "; ".join(discarded),
            )
        extracted, low, medium = _field_stats(extraction)
        usage = response.usage
        input_tokens = getattr(usage, "input_tokens", None)
        output_tokens = getattr(usage, "output_tokens", None)
        cache_read = getattr(usage, "cache_read_input_tokens", None)
        cache_creation = getattr(usage, "cache_creation_input_tokens", None)
        meta = ExtractionMeta(
            model=self._model,
            input_tokens=input_tokens,
            output_tokens=output_tokens,
            cache_read_tokens=cache_read,
            cache_creation_tokens=cache_creation,
            estimated_cost_usd=_estimate_cost(
                self._model, input_tokens, output_tokens, cache_read, cache_creation
            ),
            fields_extracted=extracted,
            low_confidence_count=low,
            medium_confidence_count=medium,
        )
        return ExtractionResult(extraction=extraction, meta=meta)


def _is_iso_date(value: str) -> bool:
    try:
        date.fromisoformat(value)
        return True
    except (ValueError, TypeError):
        return False


def _clean_property(p: ExtractedProperty) -> None:
    # getattr default None so a currently-disabled (commented-out) field is simply skipped.
    for name in ("sq_ft", "number_of_rooms", "floor", "property_tax", "house_committee"):
        v = getattr(p, name, None)
        if v is not None and v < 0:
            setattr(p, name, None)


# Security/collateral type the renter form accepts; other free text is dropped.
_SECURITY_TYPES = {"bank_guarantee", "wire_transfer"}

# Payment methods the renter form's paymentType select accepts. "wire_transfer" is a legacy
# alias the old web form stored for a bank transfer, so map it instead of dropping it.
_PAYMENT_TYPES = {"cash", "bank_transfer", "bit", "check"}
_PAYMENT_TYPE_ALIASES = {"wire_transfer": "bank_transfer"}


def _looks_like_email(value: str) -> bool:
    """Shape check only — one "@" with non-blank text either side and a dot in the domain.
    Deliberately loose: the point is to catch a phone/ID/name landing in `email`, not to
    referee which addresses are deliverable."""
    local, sep, domain = value.strip().partition("@")
    if not sep or any(c.isspace() for c in value.strip()):
        return False
    return bool(local) and "@" not in domain and "." in domain and not domain.endswith(".")


def _looks_like_israeli_id(value: str) -> bool:
    """A bare 9-digit number is almost certainly a national ID (ת\"ז), not a phone."""
    digits = "".join(ch for ch in value if ch.isdigit())
    return len(digits) == 9 and digits == value.strip()


def _clean_renter(r: ExtractedRenter, discarded: Optional[list[str]] = None) -> None:
    def _drop(field: str) -> None:
        if discarded is not None:
            discarded.append(f"{field}={getattr(r, field)!r}")
        setattr(r, field, None)

    # Mirrors RenterCreate.payment_day_in_range (1..31).
    if r.payment_day_of_month is not None and not (1 <= r.payment_day_of_month <= 31):
        _drop("payment_day_of_month")
    # lease_start must be a real ISO date; the form/back end store it as a `date`.
    if r.lease_start is not None and not _is_iso_date(r.lease_start):
        r.lease_start = None
    for name in (
        "base_rent",
        "insurance_amount",
        "rent_escalation_value",
        "contract_term_years",
        "option_years",
    ):
        v = getattr(r, name)
        if v is not None and v < 0:
            setattr(r, name, None)
    # The months are a remainder on top of the years, so anything outside 0-11 means the
    # model restated the whole term in months and the two would double-count.
    for name in ("contract_term_months", "option_term_months"):
        v = getattr(r, name)
        if v is not None and not 0 <= v <= 11:
            setattr(r, name, None)
    # Same bound per row, where a `custom` schedule states its own period lengths.
    if r.lease_years:
        for ly in r.lease_years:
            if ly.months is not None and not 1 <= ly.months <= 12:
                ly.months = None
    # Drop the whole schedule if any year's amount is negative — a partial/garbled
    # schedule is confusing to review; the user re-enters a clean one.
    if r.lease_years is not None and any(ly.amount < 0 for ly in r.lease_years):
        r.lease_years = None
    # Insurance/security type must be one of the enum values the renter form accepts.
    if r.insurance_type is not None and r.insurance_type not in _SECURITY_TYPES:
        r.insurance_type = None
    # Same for the payment method: the form's select has no option for free text, so an
    # unmapped value would render a blank control and then fail validation on submit.
    if r.payment_type is not None:
        canonical = _PAYMENT_TYPE_ALIASES.get(r.payment_type, r.payment_type)
        if canonical in _PAYMENT_TYPES:
            r.payment_type = canonical
        else:
            _drop("payment_type")
    # An email that isn't address-shaped is a phone/ID/name that wandered into the field.
    if r.email is not None and not _looks_like_email(r.email):
        _drop("email")
    # Guard against the national ID being mistaken for a phone (see the prompt).
    if r.phone is not None and _looks_like_israeli_id(r.phone):
        _drop("phone")


def _clean_extraction(
    extraction: LeaseExtraction, discarded: Optional[list[str]] = None
) -> LeaseExtraction:
    """Null out extracted values that wouldn't survive form/back-end validation, then
    drop any uncertainty notes whose field we just nulled.

    ``discarded`` is an optional out-param: each value we had to throw away is appended
    as ``"field=value"`` so the caller can log what the model actually emitted.
    """
    _clean_property(extraction.property)
    for renter in extraction.renters:
        _clean_renter(renter, discarded)
    if extraction.joint_monthly_rent is not None and extraction.joint_monthly_rent < 0:
        extraction.joint_monthly_rent = None

    def _note_target(n):
        if n.section == "property":
            return extraction.property
        if n.renter_index is not None and 0 <= n.renter_index < len(extraction.renters):
            return extraction.renters[n.renter_index]
        return None

    extraction.notes = [
        n for n in extraction.notes
        if (_t := _note_target(n)) is not None and getattr(_t, n.field, None) is not None
    ]
    return extraction


def _field_stats(extraction: LeaseExtraction) -> tuple[int, int, int]:
    """Count populated fields, and low/medium-confidence counts from the notes."""
    extracted = 0
    for section in (extraction.property, *extraction.renters):
        for value in vars(section).values():
            if value is not None:
                extracted += 1
    low = sum(1 for n in extraction.notes if n.confidence == "low")
    medium = sum(1 for n in extraction.notes if n.confidence == "medium")
    return extracted, low, medium


def _estimate_cost(
    model: str,
    input_tokens: Optional[int],
    output_tokens: Optional[int],
    cache_read: Optional[int],
    cache_creation: Optional[int],
) -> Optional[float]:
    """Rough USD estimate (cache writes ~1.25x input, cache reads ~0.1x input)."""
    price = _PRICES.get(model)
    if price is None or input_tokens is None or output_tokens is None:
        return None
    in_rate, out_rate = price
    cost = (
        input_tokens * in_rate
        + (cache_creation or 0) * in_rate * 1.25
        + (cache_read or 0) * in_rate * 0.10
        + output_tokens * out_rate
    ) / 1_000_000
    return round(cost, 6)
